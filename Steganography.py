import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from PIL import Image, ImageTk
from docx import Document
from mutagen.mp3 import MP3
from mutagen.id3 import ID3, TIT2
import re


# Function to sanitize file names by replacing invalid characters with underscores
def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', '_', filename)


# LSB steganography encoder function for images
def encode_message_image(image_path, message):
    image = Image.open(image_path)
    binary_message = ''.join(format(ord(char), '08b') for char in message)
    if len(binary_message) > (image.width * image.height * 3):
        messagebox.showerror("Error", "Message is too large for the image!")
        return None
    pixels = image.load()
    binary_message += '1111111111111110'  # End of message marker
    index = 0
    for y in range(image.height):
        for x in range(image.width):
            r, g, b = pixels[x, y]
            if index < len(binary_message):
                pixels[x, y] = (r & 254 | int(binary_message[index]), g, b)
                index += 1
            else:
                image.save("encoded_image.png")
                messagebox.showinfo("Success", "Message encoded successfully!")
                return "encoded_image.png"
    return None


# LSB steganography encoder function for Word documents
def encode_message_word(doc_path, message):
    try:
        # Open the Word document
        doc = Document(doc_path)

        # Add the message at the end of the document
        doc.add_paragraph(message)

        # Save the modified document
        encoded_doc_path = "encoded_document.docx"
        doc.save(encoded_doc_path)

        messagebox.showinfo("Success", "Message encoded successfully!")
        return encoded_doc_path
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")
        return None


# LSB steganography encoder function for audio files
def encode_message_audio(audio_path, message):
    try:
        audio = MP3(audio_path, ID3=ID3)
        if audio.tags is None:
            audio.tags = ID3()

        # Adding the text message to MP3 metadata
        audio.tags.add(TIT2(encoding=3, text=message))
        audio.save()

        messagebox.showinfo("Success", "Message encoded successfully!")
        return audio_path
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")
    return None


# LSB steganography decoder function for images
def decode_message_image(image_path):
    image = Image.open(image_path)
    pixels = image.load()
    binary_message = ''
    for y in range(image.height):
        for x in range(image.width):
            r, _, _ = pixels[x, y]
            binary_message += str(r & 1)
            if binary_message[-16:] == '1111111111111110':  # End of message marker
                message = ''.join(chr(int(binary_message[i:i + 8], 2)) for i in range(0, len(binary_message) - 16, 8))
                return message
    return None


# LSB steganography decoder function for Word documents
def decode_message_word(doc_path):
    try:
        # Open the Word document
        doc = Document(doc_path)

        # Extract the last paragraph as the hidden message
        decoded_message = doc.paragraphs[-1].text

        # Show the decoded message
        if decoded_message:
            return decoded_message
        else:
            return None
    except Exception as e:
        messagebox.showerror("Error", f"Failed to decode message from Word document: {str(e)}")
        return None


# LSB steganography decoder function for audio files
def decode_message_audio(audio_path):
    try:
        audio = MP3(audio_path, ID3=ID3)
        # Extracting text message from MP3 metadata
        message = audio.tags.get("TIT2").text[0]
        return message
    except Exception as e:
        messagebox.showerror("Error", f"Failed to decode message from audio file: {str(e)}")
        return None


# GUI
class SteganographyApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Steganography Tool")

        # Creating main frame
        self.main_frame = tk.Frame(self)
        self.main_frame.pack()

        # Adding labels and buttons
        tk.Label(self.main_frame, text="Choose an action:").pack()

        self.encode_button = tk.Button(self.main_frame, text="Encode Message", command=self.choose_encode_option)
        self.encode_button.pack(pady=5)

        self.decode_button = tk.Button(self.main_frame, text="Decode Message", command=self.choose_decode_option)
        self.decode_button.pack(pady=5)

    def choose_encode_option(self):
        encode_options = ["Encode into Image", "Encode into Word Document", "Encode into Audio"]
        option = simpledialog.askinteger("Encode Options",
                                         "Choose an option:\n1. Encode into Image\n2. Encode into Word Document\n3. Encode into Audio (MP3 format only)",
                                         minvalue=1, maxvalue=3)
        if option:
            if option == 1:
                self.encode_message_image()
            elif option == 2:
                self.encode_message_word()
            elif option == 3:
                self.encode_message_audio()

    def encode_message_image(self):
        file_path = filedialog.askopenfilename(filetypes=[("Image Files", ".png;.jpg;*.jpeg")])
        if not file_path:
            return
        message = simpledialog.askstring("Input", "Enter message to encode:")
        if message:
            encoded_image_path = encode_message_image(file_path, message)
            if encoded_image_path:
                messagebox.showinfo("Success", f"Message encoded successfully! Saved as {encoded_image_path}")

    def encode_message_word(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if not file_path:
            return
        message = simpledialog.askstring("Input", "Enter message to encode:")
        if message:
            encoded_doc_path = encode_message_word(file_path, message)
            if encoded_doc_path:
                messagebox.showinfo("Success", f"Message encoded successfully! Saved as {encoded_doc_path}")

    def encode_message_audio(self):
        file_path = filedialog.askopenfilename(filetypes=[("Audio Files", "*.mp3")])
        if not file_path:
            return
        message = simpledialog.askstring("Input", "Enter message to encode:")
        if message:
            encoded_audio_path = encode_message_audio(file_path, message)
            if encoded_audio_path:
                messagebox.showinfo("Success", f"Message encoded successfully! Saved as {encoded_audio_path}")

    def choose_decode_option(self):
        decode_options = ["Decode from Image", "Decode from Word Document", "Decode from Audio"]
        option = simpledialog.askinteger("Decode Options",
                                         "Choose an option:\n1. Decode from Image\n2. Decode from Word Document\n3. Decode from Audio (MP3 format only)",
                                         minvalue=1, maxvalue=3)
        if option:
            if option == 1:
                self.decode_message_image()
            elif option == 2:
                self.decode_message_word()
            elif option == 3:
                self.decode_message_audio()

    def decode_message_image(self):
        file_path = filedialog.askopenfilename(filetypes=[("Image Files", ".png;.jpg;*.jpeg")])
        if not file_path:
            return
        decoded_message = decode_message_image(file_path)
        if decoded_message is not None:
            messagebox.showinfo("Decoded Message", f"The decoded message is: {decoded_message}")
        else:
            messagebox.showerror("Error", "Failed to decode message from image!")

    def decode_message_word(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if not file_path:
            return
        decoded_message = decode_message_word(file_path)
        if decoded_message is not None:
            messagebox.showinfo("Decoded Message", f"The decoded message is: {decoded_message}")
        else:
            messagebox.showerror("Error", "Failed to decode message from Word document!")

    def decode_message_audio(self):
        file_path = filedialog.askopenfilename(filetypes=[("Audio Files", "*.mp3")])
        if not file_path:
            return
        decoded_message = decode_message_audio(file_path)
        if decoded_message is not None:
            messagebox.showinfo("Decoded Message", f"The decoded message is: {decoded_message}")
        else:
            messagebox.showerror("Error", "Failed to decode message from audio file!")


if __name__ == "__main__":
    app = SteganographyApp()
    app.mainloop()
